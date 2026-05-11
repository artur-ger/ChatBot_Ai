from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from typing import Any
from pypdf import PdfReader

from app.core.errors import ValidationAppError


@dataclass(frozen=True)
class ExtractedTextSegment:
    text: str
    page_number: int | None = None


_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1251", "koi8-r", "latin-1")


def detect_text_encoding(data: bytes) -> str:
    for encoding in _TEXT_ENCODINGS:
        try:
            data.decode(encoding)
        except UnicodeDecodeError:
            continue
        return encoding
    return "utf-8"


def _paragraphs_from_text(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\f", "\n\n")
    raw_blocks = re.split(r"\n\s*\n", normalized)
    paragraphs: list[str] = []
    for block in raw_blocks:
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in block.split("\n")]
        paragraph = " ".join(line for line in lines if line).strip()
        if paragraph:
            paragraphs.append(paragraph)
    return paragraphs


def _extract_text_segments_from_plain_file(path: Path) -> list[ExtractedTextSegment]:
    data = path.read_bytes()
    encoding = detect_text_encoding(data)
    text = data.decode(encoding, errors="replace")
    return [ExtractedTextSegment(text=paragraph) for paragraph in _paragraphs_from_text(text)]


def _extract_text_segments_from_pdf(path: Path) -> list[ExtractedTextSegment]:
    reader = PdfReader(str(path))
    segments: list[ExtractedTextSegment] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for paragraph in _paragraphs_from_text(text):
            segments.append(ExtractedTextSegment(text=paragraph, page_number=page_index))
    return segments


def _iter_docx_text_blocks(document: Any) -> list[str]:
    blocks: list[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(paragraph.text for paragraph in cell.paragraphs)
    return blocks


def _extract_text_segments_from_docx(path: Path) -> list[ExtractedTextSegment]:
    document = DocxDocument(str(path))
    segments: list[ExtractedTextSegment] = []
    for block in _iter_docx_text_blocks(document):
        for paragraph in _paragraphs_from_text(block):
            segments.append(ExtractedTextSegment(text=paragraph))
    return segments


def extract_text_segments_from_file(*, path: Path, doc_type: str) -> list[ExtractedTextSegment]:
    try:
        if doc_type in {"text", "markdown"}:
            return _extract_text_segments_from_plain_file(path)

        if doc_type == "pdf":
            return _extract_text_segments_from_pdf(path)

        if doc_type == "docx":
            return _extract_text_segments_from_docx(path)
    except Exception as exc:
        raise ValidationAppError(f"Failed to parse document: {exc}") from exc

    raise ValidationAppError("Unsupported document type for text extraction")


def extract_text_from_file(*, path: Path, doc_type: str) -> str:
    return "\n\n".join(
        segment.text for segment in extract_text_segments_from_file(path=path, doc_type=doc_type)
    ).strip()
