from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.core.errors import ValidationAppError
from app.services.text_chunking import chunk_text
from app.services.text_extract import extract_text_from_file


def test_extract_text_from_txt_and_markdown(tmp_path: Path) -> None:
    txt_path = tmp_path / "note.txt"
    txt_path.write_text("Привет из TXT.\n\nВторой абзац.", encoding="utf-8")
    assert "Привет из TXT" in extract_text_from_file(path=txt_path, doc_type="text")

    md_path = tmp_path / "note.md"
    md_path.write_text("# Заголовок\n\nТекст markdown.", encoding="utf-8")
    assert "markdown" in extract_text_from_file(path=md_path, doc_type="markdown")


def test_extract_text_from_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Первый параграф DOCX.")
    document.add_paragraph("Второй параграф с кириллицей.")
    document.save(str(docx_path))

    text = extract_text_from_file(path=docx_path, doc_type="docx")
    assert "Первый параграф DOCX" in text
    assert "кириллицей" in text


def test_extract_text_from_pdf(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_path.open("wb") as handle:
        writer.write(handle)

    text = extract_text_from_file(path=pdf_path, doc_type="pdf")
    assert isinstance(text, str)


def test_extract_text_unsupported_type(tmp_path: Path) -> None:
    path = tmp_path / "data.bin"
    path.write_bytes(b"123")
    with pytest.raises(ValidationAppError):
        extract_text_from_file(path=path, doc_type="unknown")


def test_chunk_text_overlap_and_dedup() -> None:
    text = "А" * 50 + "\n\n" + "Б" * 50 + "\n\n" + "А" * 50
    chunks = chunk_text(text, chunk_size=60, overlap=10)
    assert chunks
    assert len(chunks) == len(set(chunks))
    assert all(len(chunk) <= 60 for chunk in chunks)


def test_chunk_text_empty() -> None:
    assert chunk_text("   \n\n  ", chunk_size=100, overlap=10) == []
